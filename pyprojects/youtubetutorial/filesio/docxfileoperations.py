"""
    Author: Prosenjit Ghosh Chowdhury
    Date: 01-Nov-2024
"""
# Install required library/module - pip install python-docx
# How to read Word Document
# How to parse word document based on differet styles
# How to read tables from a document
# How to extract hyperlinks from a document


# Checking the working directory
import os
print(os.getcwd())

from docx.api import Document
# Read a word document
doc = Document("filesio/sampledocxread.docx")

## Get all the text ##
alltext = []
for para in doc.paragraphs:
    if para.text != '':
        alltext.append(para.text)
# Print all the text 
# print(alltext)
count = 0
for t in alltext:
    count += 1
    print(f'{count}: {t}')
    print("*****************")

## Get all the Headings ##
doc = Document("filesio/sampledocxread.docx")
headings = []
for pa in doc.paragraphs:
    if pa.style.name.startswith("Heading"):
        headings.append(pa.text)
print(headings)

## Get all the Title ##
doc = Document("filesio/sampledocxread.docx")
titles = []
for pa in doc.paragraphs:
    if pa.style.name.startswith("Title"):
        titles.append(pa.text)
print(titles)

## Get all the Keywords which are bold ##
doc = Document("filesio/sampledocxread.docx")
keywords_bold = []
for pa in doc.paragraphs:
    for run in pa.runs:
        if run.bold:
            keywords_bold.append(run.text)
print(keywords_bold)

## Get all the Keywords which are Italic ##
doc = Document("filesio/sampledocxread.docx")
keywords_italic = []
for pa in doc.paragraphs:
    for run in pa.runs:
        if run.italic:
            keywords_italic.append(run.text)
print(keywords_italic)


## Get all Highlighted text ## 
doc = Document("filesio/sampledocxread.docx")
highlighted = []
for pa in doc.paragraphs:
    for run in pa.runs:
        if run.font.highlight_color:
            highlighted.append(run.text)
print(highlighted)

## extract text from tables ##
doc = Document("filesio/sampledocxread.docx")
for table in doc.tables:
    data = [[cell.text for cell in row.cells] for row in table.rows]
print(data)
print(data[1][0])

## Extract the links form the doc ##
doc_hyperlinks = []
doc_hyperlinks_text = []
doc = Document("filesio/sampledocxread.docx")
for para in doc.paragraphs:
    for link in para.hyperlinks:
        doc_hyperlinks_text.append(link.text)
        doc_hyperlinks.append(link.address)
link_dict = {key: val for key, val in zip(doc_hyperlinks_text, doc_hyperlinks)}
print(link_dict)